"""Export service for generating various document formats."""
import logging
import io
import json
import zipfile
from datetime import datetime
from typing import BinaryIO, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

from ..models import KPISchema
from .bpmn_processor import BPMNProcessor

logger = logging.getLogger(__name__)


class ExportService:
    """Service for exporting process data in various formats."""

    def __init__(self):
        """Initialize export service."""
        self.bpmn_processor = BPMNProcessor()

    def export_to_docx(
        self,
        dpt: dict,
        bpmn_json: Optional[dict] = None,
        kpis: Optional[list] = None,
        filename: str = "ceproc_export.docx"
    ) -> BinaryIO:
        """
        Export process data to Word document.

        Args:
            dpt: DPT dictionary with process information
            bpmn_json: Optional BPMN JSON for diagram
            kpis: Optional list of KPI indicators
            filename: Output filename

        Returns:
            BytesIO object with DOCX content
        """
        logger.info(f"Generating DOCX export: {filename}")

        doc = Document()

        # Add title
        title = doc.add_heading("Análise de Processo", level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add metadata
        metadata = dpt.get("metadados", {})
        doc.add_heading("1. Informações Gerais", level=2)
        doc.add_paragraph(f"Processo: {metadata.get('processo', 'N/A')}")
        doc.add_paragraph(f"Data de Análise: {metadata.get('data_analise', 'N/A')}")
        doc.add_paragraph(f"Analista: {metadata.get('analista', 'N/A')}")
        doc.add_paragraph(f"Departamento: {metadata.get('departamento', 'N/A')}")

        # Add business context
        negocio = dpt.get("negocio", {})
        doc.add_heading("2. Contexto de Negócio", level=2)
        doc.add_paragraph(f"Descrição: {negocio.get('descricao', 'N/A')}")
        doc.add_paragraph(f"Objetivos: {negocio.get('objetivos', 'N/A')}")
        doc.add_paragraph(f"Restrições: {negocio.get('restricoes', 'N/A')}")

        # Add purpose
        finalidade = dpt.get("finalidade", {})
        doc.add_heading("3. Finalidade do Processo", level=2)
        doc.add_paragraph(f"Objetivo Principal: {finalidade.get('objetivo_principal', 'N/A')}")
        doc.add_paragraph(f"Escopo: {finalidade.get('escopo', 'N/A')}")

        # Add process flow
        doc.add_heading("4. Fluxo de Processo", level=2)
        etapas = dpt.get("etapas", [])
        if etapas:
            for i, etapa in enumerate(etapas, 1):
                p = doc.add_paragraph(style='List Number')
                p.add_run(f"{etapa.get('titulo', f'Etapa {i}')}").bold = True
                p.add_run(f": {etapa.get('descricao', 'N/A')}")

        # Add actors
        doc.add_heading("5. Atores e Responsabilidades", level=2)
        atores = dpt.get("atores", [])
        if atores:
            for ator in atores:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{ator.get('nome', 'N/A')}").bold = True
                p.add_run(f" - {ator.get('responsabilidades', 'N/A')}")

        # Add systems
        doc.add_heading("6. Sistemas Envolvidos", level=2)
        sistemas = dpt.get("sistemas", [])
        if sistemas:
            for sistema in sistemas:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{sistema.get('nome', 'N/A')}").bold = True
                p.add_run(f": {sistema.get('descricao', 'N/A')}")

        # Add KPIs
        if kpis:
            doc.add_heading("7. Indicadores de Desempenho", level=2)
            table = doc.add_table(rows=1, cols=7)
            table.style = 'Light Grid Accent 1'

            # Header row
            header_cells = table.rows[0].cells
            headers = ["Indicador", "Objetivo", "Unidade", "Meta", "Periodicidade", "Criticidade", "Responsável"]
            for i, header in enumerate(headers):
                header_cells[i].text = header

            # Data rows
            for kpi in kpis:
                row_cells = table.add_row().cells
                row_cells[0].text = str(kpi.get("indicador", ""))
                row_cells[1].text = str(kpi.get("objetivo", ""))
                row_cells[2].text = str(kpi.get("unidade", ""))
                row_cells[3].text = str(kpi.get("meta", ""))
                row_cells[4].text = str(kpi.get("periodicidade", ""))
                row_cells[5].text = str(kpi.get("criticidade", ""))
                row_cells[6].text = str(kpi.get("responsavel", ""))

        # Add footer with generation timestamp
        doc.add_paragraph()
        footer = doc.add_paragraph(
            f"Documento gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M:%S')}"
        )
        footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        footer_format = footer.runs[0]
        footer_format.font.size = Pt(10)
        footer_format.font.color.rgb = RGBColor(128, 128, 128)

        # Save to BytesIO
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        logger.info("DOCX export completed successfully")
        return output

    def export_to_xlsx(
        self,
        kpis: list,
        filename: str = "ceproc_kpis.xlsx"
    ) -> BinaryIO:
        """
        Export KPI indicators to Excel spreadsheet.

        Args:
            kpis: List of KPI dictionaries
            filename: Output filename

        Returns:
            BytesIO object with XLSX content
        """
        logger.info(f"Generating XLSX export: {filename}")

        wb = Workbook()
        ws = wb.active
        ws.title = "KPIs"

        # Define headers
        headers = [
            "Indicador",
            "Objetivo",
            "Processo",
            "Cliente",
            "Metadados",
            "Fonte de Extração",
            "Fórmula de Cálculo",
            "Unidade",
            "Filtro",
            "Meta",
            "Periodicidade",
            "Polaridade",
            "Responsável",
            "Criticidade",
            "Justificativa"
        ]

        # Add headers to first row
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col)
            cell.value = header
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="1a3a52", end_color="1a3a52", fill_type="solid")
            cell.alignment = Alignment(horizontal="center", vertical="center")

        # Add KPI data rows
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        for row, kpi in enumerate(kpis, 2):
            ws.cell(row=row, column=1).value = kpi.get("indicador", "")
            ws.cell(row=row, column=2).value = kpi.get("objetivo", "")
            ws.cell(row=row, column=3).value = kpi.get("processo", "")
            ws.cell(row=row, column=4).value = kpi.get("cliente", "")
            ws.cell(row=row, column=5).value = kpi.get("metadados", "")
            ws.cell(row=row, column=6).value = kpi.get("fonte_extracao", "")
            ws.cell(row=row, column=7).value = kpi.get("formula_calculo", "")
            ws.cell(row=row, column=8).value = kpi.get("unidade", "")
            ws.cell(row=row, column=9).value = kpi.get("filtro", "")
            ws.cell(row=row, column=10).value = kpi.get("meta", "")
            ws.cell(row=row, column=11).value = kpi.get("periodicidade", "")
            ws.cell(row=row, column=12).value = kpi.get("polaridade", "")
            ws.cell(row=row, column=13).value = kpi.get("responsavel", "")
            ws.cell(row=row, column=14).value = kpi.get("criticidade", "")
            ws.cell(row=row, column=15).value = kpi.get("justificativa", "")

            # Apply borders to all cells
            for col in range(1, len(headers) + 1):
                ws.cell(row=row, column=col).border = border
                ws.cell(row=row, column=col).alignment = Alignment(wrap_text=True, vertical="top")

        # Set column widths
        column_widths = [20, 15, 15, 15, 15, 20, 25, 10, 15, 10, 15, 12, 15, 12, 25]
        for col, width in enumerate(column_widths, 1):
            ws.column_dimensions[chr(64 + col)].width = width

        # Save to BytesIO
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        logger.info("XLSX export completed successfully")
        return output

    def export_to_bpmn_xml(
        self,
        bpmn_json: dict,
        filename: str = "process.bpmn"
    ) -> BinaryIO:
        """
        Export BPMN JSON to BPMN 2.0 XML format.

        Args:
            bpmn_json: BPMN structure from LLM
            filename: Output filename

        Returns:
            BytesIO object with BPMN XML content
        """
        logger.info(f"Generating BPMN XML export: {filename}")

        try:
            xml_content = self.bpmn_processor.convert_json_to_xml(bpmn_json)

            output = io.BytesIO()
            output.write(xml_content.encode('utf-8'))
            output.seek(0)

            logger.info("BPMN XML export completed successfully")
            return output

        except Exception as e:
            logger.error(f"Error exporting to BPMN XML: {str(e)}")
            raise

    def export_to_zip(
        self,
        dpt: dict,
        bpmn_json: Optional[dict] = None,
        bpmn_xml: Optional[str] = None,
        kpis: Optional[list] = None,
        filename: str = "ceproc_export.zip"
    ) -> BinaryIO:
        """
        Export all data as ZIP archive containing multiple formats.

        Args:
            dpt: DPT dictionary
            bpmn_json: Optional BPMN JSON
            bpmn_xml: Optional BPMN XML string
            kpis: Optional KPI list
            filename: Output filename

        Returns:
            BytesIO object with ZIP content
        """
        logger.info(f"Generating ZIP export: {filename}")

        output = io.BytesIO()

        with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as zipf:
            # Add DPT JSON
            dpt_json = json.dumps(dpt, indent=2, ensure_ascii=False)
            zipf.writestr("dpt.json", dpt_json.encode('utf-8'))

            # Add BPMN JSON if available
            if bpmn_json:
                bpmn_json_str = json.dumps(bpmn_json, indent=2, ensure_ascii=False)
                zipf.writestr("bpmn.json", bpmn_json_str.encode('utf-8'))

            # Add BPMN XML if available
            if bpmn_xml:
                zipf.writestr("process.bpmn", bpmn_xml.encode('utf-8'))

            # Add KPIs as JSON
            if kpis:
                kpis_json = json.dumps(kpis, indent=2, ensure_ascii=False)
                zipf.writestr("kpis.json", kpis_json.encode('utf-8'))

                # Add KPIs as XLSX
                xlsx_buffer = self.export_to_xlsx(kpis)
                zipf.writestr("kpis.xlsx", xlsx_buffer.getvalue())

            # Add DOCX report
            docx_buffer = self.export_to_docx(dpt, bpmn_json, kpis)
            zipf.writestr("relatorio.docx", docx_buffer.getvalue())

        output.seek(0)

        logger.info("ZIP export completed successfully")
        return output

    def generate_export_filename(self, export_format: str, process_name: str = "") -> str:
        """
        Generate filename for export.

        Args:
            export_format: Format type (docx, xlsx, bpmn, zip)
            process_name: Optional process name

        Returns:
            Generated filename with timestamp
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_name = process_name.replace(" ", "_") if process_name else "ceproc_export"

        extensions = {
            "docx": "docx",
            "xlsx": "xlsx",
            "bpmn": "bpmn",
            "zip": "zip"
        }

        extension = extensions.get(export_format, "zip")

        return f"{base_name}_{timestamp}.{extension}"


def get_export_service() -> ExportService:
    """Get export service instance."""
    return ExportService()
